from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "Cricket Verse Project Documentation.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(10)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_h(doc: Document, text: str, level: int = 1):
    return doc.add_heading(text, level=level)


def add_p(doc: Document, text: str = "", style: str | None = None):
    return doc.add_paragraph(text, style=style)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        add_p(doc, item, "List Bullet")


def add_numbers(doc: Document, items: list[str]) -> None:
    for item in items:
        add_p(doc, item, "List Number")


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_text(hdr[idx], header, True)
        set_cell_shading(hdr[idx], "E8EEF5")
        hdr[idx].width = Inches(widths[idx])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
            cells[idx].width = Inches(widths[idx])
    doc.add_paragraph()


def add_code_block(doc: Document, lines: list[str]) -> None:
    para = doc.add_paragraph()
    for line in lines:
        run = para.add_run(line + "\n")
        run.font.name = "Consolas"
        run.font.size = Pt(9)


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.2

    for name, size, color in [
        ("Title", 22, "0B2545"),
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)


def build() -> None:
    doc = Document()
    style_doc(doc)

    title = doc.add_paragraph()
    title.style = "Title"
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Cricket Verse Telegram Bot - Project Documentation").bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(
        "Updated local guide for match rules, DRS, AI commands, match history, virtual credits, and Render deployment."
    )

    add_h(doc, "1. Current Project Snapshot")
    add_bullets(
        doc,
        [
            "Project folder: C:\\Users\\tyson\\OneDrive\\Documents\\Cricket Bot 2",
            "Main entry point: main.py",
            "Telegram framework: python-telegram-bot with job queue and webhooks extras",
            "Database: SQLite by default, DATABASE_PATH controls the file location",
            "AI model: gemini-flash-latest through GEMINI_MODEL",
            "Render free deploy: webhook mode with RUN_MODE=webhook",
            "This document now records the May 27, 2026 Cricket Verse rules update.",
        ],
    )

    add_h(doc, "2. Latest Rule Update")
    add_bullets(
        doc,
        [
            "Natural AI replies are removed. The bot answers match questions only through /ask.",
            "/buzz is reserved for previous-match history, player stats, records, and database-backed banter.",
            "Every completed match gets a match id. /matchin <id> shows saved match details and all player data.",
            "Each team receives 2 DRS reviews per innings for LBW or Stumped wicket calls.",
            "Player of the Match is announced at match end using batting, bowling, fielding, and winning impact.",
            "Team formation should avoid chat spam by editing the same /myteam message for add, remove, and select flows.",
            "Virtual credits and fun games are for entertainment only. No real money, deposits, withdrawals, or cash value.",
        ],
    )

    add_h(doc, "3. Commands And Lobby Flow")
    add_numbers(
        doc,
        [
            "Captain sends /playmatch <overs> in the group.",
            "The bot asks both captains to join on the main match message.",
            "Only the requested captain's reply is accepted when the bot asks for team names.",
            "The bot randomly selects one captain as toss winner.",
            "Toss winner chooses Bat or Bowl.",
            "Captains use /myteam to manage their own team only.",
            "Start stays locked until both teams have equal players and the opening batter and bowler are selected.",
            "After both captains press Start, admin approval is required before the match begins.",
        ],
    )

    add_h(doc, "4. Team Formation")
    add_bullets(
        doc,
        [
            "/myteam is captain-only and displays only that captain's roster.",
            "The /myteam Add, Remove, and Select flows should edit the same menu message instead of sending many new messages.",
            "/add lets a captain add a replied player, tagged player, username, Telegram link, numeric id, or known user.",
            "Remove remains serial-number based from the captain's own list.",
            "Select shows serial buttons for choosing the opening batter or current bowler.",
            "Players are shown by Telegram first name when available. Names refresh when the bot sees the user again.",
            "Button permissions stay strict. If the wrong person clicks, the popup says: You are not [Player Name]!",
        ],
    )

    add_h(doc, "5. Live Arena MPMM UI")
    add_bullets(
        doc,
        [
            "The Main Playing Match Message updates every ball with runs, wickets, overs, batter, bowler, and last delivery.",
            "CRR always appears. RRR appears only in the second innings.",
            "Current batter shows runs, balls, and strike rate.",
            "Current bowler shows wickets, overs, runs conceded, and economy.",
            "Timeline shows the last 8 events such as 1 | 4 | wd | W(c).",
            "Extra balls receive normal short commentary, not giant AI paragraphs.",
            "Wicket, catch, drop, DRS, chase, and pressure events get human-style funny commentary.",
        ],
    )

    add_h(doc, "6. Pacer Ball Data")
    add_table(
        doc,
        ["No.", "Delivery", "Allowed Lengths", "MLR"],
        [
            ["0", "Reverse Swing", "Yorker, Full", "0, 1"],
            ["1", "Bouncer", "Bouncer, Short", "0, 1, 2"],
            ["2", "Yorker", "Yorker", "0, 1, 2"],
            ["3", "Short", "Short", "1, 2, 3, 4"],
            ["4", "Slower", "Full, Good", "Full: 1, 2, 4; Good: 0, 1, 2, 3"],
            ["6", "Knuckle", "Full, Yorker", "Full: 2, 4, 6; Yorker: 0, 2"],
        ],
        [0.6, 1.5, 1.7, 2.7],
    )

    add_h(doc, "7. Spinner Ball Data")
    add_table(
        doc,
        ["No.", "Delivery", "Allowed Lengths", "MLR"],
        [
            ["0", "Carrom", "Good", "0, 1"],
            ["1", "Doosra", "Good", "0, 1, 2"],
            ["2", "Leg Break", "Full, Good", "Full: 1, 2, 4; Good: 1, 2, 3"],
            ["3", "Top Spin", "Good, Short", "Good: 0, 1, 2, 3; Short: 1, 2, 3, 4"],
            ["4", "Flipper", "Full, Good", "Full: 2, 4, 6; Good: 1, 2, 4"],
            ["6", "Googly", "Good, Full", "Good: 0, 1, 2; Full: 2, 4, 6"],
        ],
        [0.6, 1.5, 1.7, 2.7],
    )

    add_h(doc, "8. Scoring And Hard Balls")
    add_numbers(
        doc,
        [
            "Bowler chooses a delivery type.",
            "Bot secretly chooses an actual length from that delivery's allowed lengths.",
            "Batter chooses a run value.",
            "Batter chooses a length guess.",
            "Length match gives the batter's selected runs.",
            "Length mismatch gives miss-length runs by the delivery MLR data and weighted miss-run table.",
            "Normal balls only give batter-selected runs on a length match.",
            "Hard balls plus bouncers share one cap of 3 hard-ball slots per over.",
            "For the first 3 hard-ball slots, batter gets selected runs only when length matches; otherwise miss-length runs are added.",
            "If 2 bouncers are used in an over, only 1 more hard-ball slot remains.",
            "From the 4th hard ball onward, the batter gets selected runs regardless of length.",
            "Bowler bouncer delivery is locked after 2 bouncers in the over.",
            "The batter-side bouncer length option is single-use per over if it is offered; after one use it is removed for that over.",
        ],
    )
    add_table(
        doc,
        ["Style", "Hard Balls", "Catch Balls"],
        [
            ["Pacer", "Reverse Swing (0), Bouncer (1), Yorker (2)", "Bouncer, Short, Slower, Knuckle"],
            ["Spinner", "Carrom (0), Doosra (1), Top Spin (3)", "Doosra, Leg Break, Top Spin, Googly"],
        ],
        [1.2, 2.7, 2.6],
    )
    add_table(
        doc,
        ["Miss-Length Run", "Probability"],
        [
            ["6", "10%"],
            ["4", "15%"],
            ["3", "20%"],
            ["2", "30%"],
            ["1", "25%"],
        ],
        [2.5, 2.0],
    )

    add_h(doc, "9. Wickets, Catch, Run Out, And DRS")
    add_table(
        doc,
        ["Condition", "Length Result", "Outcome Probabilities"],
        [
            ["BFR = BR", "Mismatch", "50% Bowled/Stumped/LBW; 30% Catch Out type 1; 20% Run Out"],
            ["BFR = BR", "Match", "30% Bowled/Stumped/LBW; 70% survives with batter runs"],
            ["BFR != BR", "Mismatch", "70% gets runs by ball data; 30% Catch Out type 2 only on catch-ball deliveries"],
            ["BFR != BR", "Match", "Runs proceed; type 2 catch does not trigger"],
        ],
        [1.3, 1.2, 4.0],
    )
    add_bullets(
        doc,
        [
            "Run out is a direct wicket when triggered.",
            "Catch system selects a random fielder from the bowling team and gives 120, 150, or 180 seconds.",
            "Catch time probability: 120s = 30%, 150s = 50%, 180s = 20%.",
            "Fielder guesses 0, 1, 2, 3, 4, or 6. Correct guess is OUT; wrong guess or timeout is DROP CATCH.",
            "Drop-catch bonus: 120s adds 1 run, 150s adds 2 runs, 180s adds 3 runs.",
            "Each team gets 2 DRS reviews per innings.",
            "DRS appears only on LBW or Stumped wicket commentary messages.",
            "Only the respected team captain can press DRS.",
            "DRS result: 70% wicket stays and review is lost; 30% batter survives and returns to the pitch.",
        ],
    )

    add_h(doc, "10. Extras And Spam Protocol")
    add_table(
        doc,
        ["Pacer Delivery", "Wide", "No Ball", "Leg Bye", "Spam Wide", "Spam No Ball"],
        [
            ["Reverse Swing", "4%", "1%", "2%", "75%", "20%"],
            ["Bouncer", "8%", "6%", "1%", "100%", "60%"],
            ["Yorker", "5%", "3%", "2%", "90%", "50%"],
            ["Short", "6%", "2%", "1%", "80%", "20%"],
            ["Slower", "4%", "2%", "1%", "65%", "15%"],
            ["Knuckle", "3%", "5%", "1%", "50%", "45%"],
        ],
        [1.45, 0.9, 0.9, 0.9, 1.05, 1.15],
    )
    add_table(
        doc,
        ["Spinner Delivery", "Wide", "No Ball", "Leg Bye", "Spam Wide", "Spam No Ball"],
        [
            ["Carrom", "3%", "1%", "1%", "60%", "10%"],
            ["Doosra", "5%", "1%", "1%", "80%", "15%"],
            ["Leg Break", "6%", "1%", "2%", "90%", "10%"],
            ["Top Spin", "4%", "2%", "1%", "75%", "20%"],
            ["Flipper", "5%", "3%", "1%", "70%", "25%"],
            ["Googly", "7%", "2%", "1%", "100%", "15%"],
        ],
        [1.45, 0.9, 0.9, 0.9, 1.05, 1.15],
    )
    add_table(
        doc,
        ["Spam Trigger", "Result"],
        [
            ["2nd consecutive same delivery", "20% Wide, 80% legal"],
            ["3rd consecutive same delivery", "70% No Ball, 30% legal"],
            ["4th use of same delivery within an over", "80% No Ball, 20% Wide"],
        ],
        [2.5, 4.0],
    )

    add_h(doc, "11. Innings, Match ID, And Finish")
    add_bullets(
        doc,
        [
            "First innings ends when overs finish or all available batters are out.",
            "A full batting and bowling scorecard is sent after innings completion.",
            "Second innings swaps roles and target is first innings score plus 1.",
            "Chasing team wins immediately after reaching target. Otherwise the defending team wins by runs.",
            "Every completed match is saved with a unique match id.",
            "/matchin <id> shows the saved match summary, innings scorecards, player stats, key moments, and player of the match.",
            "Player of the Match should be calculated from runs, wickets, catches, run outs, clutch chase/defense impact, and winning contribution.",
        ],
    )

    add_h(doc, "12. AI Commands")
    add_table(
        doc,
        ["Command", "Purpose", "Style"],
        [
            ["/ask <question>", "Answers only ongoing-match questions such as who may win, where the match turned, who choked, and what the next pressure point is.", "Short, funny, English-only, match-related, roast/bully energy without unrelated trash."],
            ["/buzz <question>", "Uses database history for previous match details, player stats, top run scorers, wicket leaders, best matches, and player comparisons.", "Human-style cricket banter with useful facts first and jokes second."],
            ["/matchin <id>", "Shows one saved match by id with every important score and player detail.", "Readable summary, not a raw database dump."],
            ["/myprofile", "Shows the user's profile, virtual credits, game record, and cricket stats.", "Compact and playful."],
        ],
        [1.3, 3.1, 2.1],
    )
    add_bullets(
        doc,
        [
            "The bot should not auto-reply to normal tagged/replied questions anymore.",
            "AI answers must use current match state for /ask and persisted database records for /buzz.",
            "Avoid huge paragraphs. Give the useful answer, then a quick human-style comment.",
            "Gemini model remains gemini-flash-latest, with local fallback answers if GEMINI_API_KEY is missing.",
        ],
    )

    add_h(doc, "13. Virtual Credits And Fun Games")
    add_bullets(
        doc,
        [
            "All money/coins are virtual in-bot credits only. They have no real-world value.",
            "No deposits, no withdrawals, no buying credits, and no cash prizes.",
            "/myprofile shows credits, match stats, mini-game wins, and losses.",
            "PvP games can include Toss Duel, Number Smash, Run Race, Wicket Hunt, and Over Predictor.",
            "Group games can include prediction pools using virtual credits, with the winning players sharing the virtual prize pot.",
            "Game messages should be short, funny, and safe for group play.",
        ],
    )

    add_h(doc, "14. Database And Persistence")
    add_bullets(
        doc,
        [
            "SQLite stores live match snapshots, completed matches, career stats, player names/usernames, and virtual credit profiles.",
            "Active match data should be saved after each important state change.",
            "Completed match snapshots should preserve enough ball-by-ball and player detail for /matchin, /ask, and /buzz.",
            "Free Render storage is temporary unless a persistent disk or external database is added later.",
        ],
    )

    add_h(doc, "15. Render Free Deployment")
    add_p(doc, "Use Render Web Service on the Free plan. The bot runs in webhook mode.")
    add_table(
        doc,
        ["Setting", "Value"],
        [
            ["Build Command", "pip install -r requirements.txt"],
            ["Start Command", "python main.py"],
            ["RUN_MODE", "webhook"],
            ["WEBHOOK_PATH", "telegram-webhook"],
            ["PYTHON_VERSION", "3.12.8"],
            ["GEMINI_MODEL", "gemini-flash-latest"],
            ["DATABASE_PATH", "/tmp/cricket_verse.sqlite3"],
        ],
        [2.0, 4.5],
    )
    add_p(doc, "Add TELEGRAM_BOT_TOKEN and GEMINI_API_KEY manually in Render Environment. Do not commit .env.")
    add_p(doc, "For long-term stats on Render Free, move DATABASE_PATH to a persistent disk or external database later.")

    add_h(doc, "16. Updated Files")
    add_table(
        doc,
        ["File", "Purpose", "Latest Spec Coverage"],
        [
            ["main.py", "Starts the bot", "Calls cricket_verse.bot.run()."],
            ["cricket_verse/bot.py", "Telegram commands, callback buttons, lobby, gameplay flow", "Needs command surface for /ask, /buzz, /matchin, /myprofile, DRS buttons, same-message /myteam menus, and virtual games."],
            ["cricket_verse/engine.py", "Cricket scoring and wicket engine", "Needs latest hard-ball, wicket probability, spam percentage, DRS, and player-of-match logic."],
            ["cricket_verse/game_data.py", "Pacer/spinner delivery tables", "Stores delivery lengths, MLR, hard balls, catch balls, extras, and spam data."],
            ["cricket_verse/formatting.py", "MPMM and scorecard text", "Needs compact human UI, player of match, /matchin summary, and profile formatting."],
            ["cricket_verse/database.py", "SQLite persistence", "Needs match ids, completed-match query helpers, virtual credits, and richer history for /buzz."],
            ["cricket_verse/gemini.py", "Gemini-backed answers", "Should answer only /ask and /buzz, English-only, short, funny, database-backed."],
            ["cricket_verse/config.py", "Environment settings", "RUN_MODE, PORT, webhook path/url, Gemini model, database path."],
            ["render.yaml", "Render deployment blueprint", "Free web service, Python 3.12.8, webhook env setup."],
            [".env.example", "Local env template", "Uses gemini-flash-latest and webhook/polling switches."],
            ["README.md", "Human quick guide", "Should mirror this command and rule update after implementation."],
        ],
        [1.7, 2.0, 2.8],
    )

    add_h(doc, "17. Important Environment Variables")
    add_code_block(
        doc,
        [
            "TELEGRAM_BOT_TOKEN=your_real_bot_token",
            "GEMINI_API_KEY=your_real_gemini_key",
            "GEMINI_MODEL=gemini-flash-latest",
            "RUN_MODE=webhook",
            "WEBHOOK_PATH=telegram-webhook",
            "DATABASE_PATH=/tmp/cricket_verse.sqlite3",
            "PYTHON_VERSION=3.12.8",
        ],
    )

    add_h(doc, "18. Current Function Map")
    add_table(
        doc,
        ["Module", "Key Functions"],
        [
            ["bot.py", "playmatch, myteam, callbacks, handle_bowl, handle_bat, handle_length, handle_ready, handle_catch, event_commentary, run"],
            ["engine.py", "choose_length, make_pending_delivery, resolve_pending_delivery, score_runs_for_ball, wicket_check, finish_catch"],
            ["formatting.py", "scoreboard, team_roster_text, teams_text, innings_scorecard, match_summary"],
            ["database.py", "save_match, load_match, complete_match, apply_match_stats, player_stats, player_by_username"],
            ["gemini.py", "summarize_player, answer_player_question, local_summary, local_question_answer"],
            ["models.py", "Match, make_match, reset_score, new_player, new_player_stats, user_label"],
            ["game_data.py", "Delivery, PACER_DELIVERIES, SPIN_DELIVERIES, delivery_label"],
        ],
        [1.5, 5.0],
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_h(doc, "19. Final Notes")
    add_bullets(
        doc,
        [
            "This document is the updated local Cricket Verse specification for the next code pass.",
            "Render should be redeployed after GitHub receives the updated files.",
            "The local machine does not currently expose a working git command in PowerShell, so GitHub Desktop or website upload is the easiest route.",
        ],
    )

    doc.save(OUT)


if __name__ == "__main__":
    build()
    print(OUT)
